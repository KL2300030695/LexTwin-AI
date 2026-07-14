from docx import Document
from docx.enum.text import WD_BREAK

from app.models.schema import DocType
from app.parsers.docx_parser import parse_docx


def _make_sample_docx(path: str) -> None:
    doc = Document()
    doc.add_paragraph("SAMPLE SERVICES AGREEMENT")
    doc.add_paragraph("4 PAYMENT TERMS")
    doc.add_paragraph("Provider shall invoice monthly as described below.")
    doc.add_paragraph("4.1 Invoicing")
    doc.add_paragraph("Client shall pay within Section 4.2 of this Agreement.")
    doc.add_paragraph("4.2 Late Fees")
    doc.add_paragraph("Notwithstanding Section 4.1, late fees accrue after 10 days.")

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Days Late"
    table.rows[0].cells[1].text = "Fee"
    table.rows[1].cells[0].text = "1-10"
    table.rows[1].cells[1].text = "1%"

    doc.add_paragraph("4.3 Refer to Exhibit D for the fee schedule.")

    # Explicit page break, followed by a second-page heading
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

    doc.add_paragraph("5 LIABILITY")
    doc.add_paragraph("Liability is capped as described herein.")

    doc.save(path)


def test_docx_parses_hierarchy_tables_and_references(tmp_path):
    path = str(tmp_path / "sample.docx")
    _make_sample_docx(path)

    parsed = parse_docx(path, "doc1", "sample.docx", DocType.MSA)
    by_num = {c.section_number: c for c in parsed.clauses}

    assert by_num["4.1"].parent_section == "4"
    assert by_num["4.2"].parent_section == "4"
    assert "Section 4.2" in by_num["4.1"].text

    nw_refs = [r for r in by_num["4.2"].references if r.is_notwithstanding]
    assert len(nw_refs) == 1
    assert nw_refs[0].target_section == "4.1"

    assert len(parsed.tables) == 1
    assert parsed.tables[0].rows == [["Days Late", "Fee"], ["1-10", "1%"]]
    assert parsed.tables[0].id in by_num["4.2"].table_ids

    assert "Exhibit D" in parsed.known_exhibit_labels
    assert "Exhibit D" not in parsed.available_exhibit_labels


def test_docx_page_break_increments_page_number(tmp_path):
    path = str(tmp_path / "sample.docx")
    _make_sample_docx(path)

    parsed = parse_docx(path, "doc1", "sample.docx", DocType.MSA)
    by_num = {c.section_number: c for c in parsed.clauses}

    assert by_num["4"].page_start == 1
    assert by_num["5"].page_start == 2
    assert parsed.page_count == 2
