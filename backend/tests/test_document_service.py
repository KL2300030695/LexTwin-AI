from docx import Document

from app.models.schema import DocType
from app.services.document_service import get_document, ingest_document, replace_document


def _make_docx(path: str, heading: str, body: str) -> None:
    doc = Document()
    doc.add_paragraph(heading)
    doc.add_paragraph("1 FIRST CLAUSE")
    doc.add_paragraph(body)
    doc.save(path)


def test_ingest_document_generates_a_fresh_doc_id_each_time(tmp_path):
    path = str(tmp_path / "v1.docx")
    _make_docx(path, "SAMPLE AGREEMENT", "Provider shall deliver within thirty days.")

    first = ingest_document(path, "sample.docx", DocType.MSA)
    second = ingest_document(path, "sample.docx", DocType.MSA)

    assert first.doc_id != second.doc_id
    assert first.doc_id.startswith("msa-")
    assert get_document(first.doc_id) is not None
    assert get_document(second.doc_id) is not None


def test_replace_document_keeps_the_same_doc_id_and_updates_content(tmp_path):
    v1_path = str(tmp_path / "v1.docx")
    _make_docx(v1_path, "SAMPLE AGREEMENT", "Provider shall deliver within thirty days.")
    original = ingest_document(v1_path, "sample.docx", DocType.MSA)

    v2_path = str(tmp_path / "v2.docx")
    _make_docx(v2_path, "SAMPLE AGREEMENT", "Provider shall deliver within sixty days.")
    updated = replace_document(original.doc_id, v2_path, "sample_v2.docx", DocType.MSA)

    assert updated.doc_id == original.doc_id
    assert any("sixty days" in c.text for c in updated.clauses)

    refetched = get_document(original.doc_id)
    assert refetched is not None
    assert any("sixty days" in c.text for c in refetched.clauses)
    assert refetched.filename == "sample_v2.docx"
