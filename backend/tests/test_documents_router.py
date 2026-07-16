"""End-to-end check of /api/documents (upload, get, list, replace) against
the real FastAPI app, over real multipart requests -- no mocking, since
document ingestion is the one place a raw file actually needs to round-trip
through parsing."""
from docx import Document
from fastapi.testclient import TestClient

from app.main import app

client = TestClient(app)


def _docx_bytes(body: str) -> bytes:
    import io

    doc = Document()
    doc.add_paragraph("SAMPLE AGREEMENT")
    doc.add_paragraph("1 FIRST CLAUSE")
    doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_upload_then_replace_keeps_the_same_doc_id():
    upload_resp = client.post(
        "/api/documents/upload",
        files={"file": ("v1.docx", _docx_bytes("Provider shall deliver within thirty days."), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"doc_type": "MSA"},
    )
    assert upload_resp.status_code == 200
    doc_id = upload_resp.json()["doc_id"]

    replace_resp = client.put(
        f"/api/documents/{doc_id}/replace",
        files={"file": ("v2.docx", _docx_bytes("Provider shall deliver within sixty days."), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert replace_resp.status_code == 200
    body = replace_resp.json()
    assert body["doc_id"] == doc_id
    assert any("sixty days" in c["text"] for c in body["clauses"])

    get_resp = client.get(f"/api/documents/{doc_id}")
    assert get_resp.status_code == 200
    assert any("sixty days" in c["text"] for c in get_resp.json()["clauses"])

    list_resp = client.get("/api/documents")
    assert list_resp.status_code == 200
    assert sum(1 for d in list_resp.json() if d["doc_id"] == doc_id) == 1


def test_replace_unknown_doc_id_returns_404():
    resp = client.put(
        "/api/documents/msa-doesnotexist/replace",
        files={"file": ("v2.docx", _docx_bytes("Provider shall deliver within sixty days."), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 404
